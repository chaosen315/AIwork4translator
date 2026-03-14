import json
import os
import re
import html
from typing import List, Dict, Any
import pandas as pd
from docx import Document
from docx.shared import RGBColor
from xhtml2pdf import pisa
from ebooklib import epub

def _load_json_data(json_path: str) -> List[Dict[str, Any]]:
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    return data.get('text_info', [])

def _iter_markdown_blocks(text: str):
    lines = (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    buffer: List[str] = []
    for line in lines:
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if m:
            paragraph_text = "\n".join(buffer).strip()
            if paragraph_text:
                for p in paragraph_text.split("\n\n"):
                    p = p.strip()
                    if p:
                        yield ("para", p)
            buffer = []
            yield ("heading", len(m.group(1)), m.group(2).strip())
        else:
            buffer.append(line)

    paragraph_text = "\n".join(buffer).strip()
    if paragraph_text:
        for p in paragraph_text.split("\n\n"):
            p = p.strip()
            if p:
                yield ("para", p)

def export_json_to_xlsx(json_path: str, output_path: str) -> bool:
    try:
        data = _load_json_data(json_path)
        rows = []
        for item in data:
            rows.append({
                'ID': item.get('paragraph_number'),
                'Source': item.get('content'),
                'Translation': item.get('translation'),
                'Notes': item.get('notes'),
                'Header Path': " > ".join(item.get('meta_data', {}).get('header_path', []) or []) if item.get('meta_data') else ""
            })
        
        df = pd.DataFrame(rows)
        df.to_excel(output_path, index=False)
        return True
    except Exception as e:
        print(f"[Error] Export to XLSX failed: {e}")
        return False

def export_json_to_docx(json_path: str, output_path: str) -> bool:
    try:
        data = _load_json_data(json_path)
        doc = Document()
        
        # Add Title
        base_name = os.path.splitext(os.path.basename(json_path))[0].replace('_intermediate', '')
        doc.add_heading(base_name, 0)
        
        for item in data:
            content = item.get('translation', '') or ''
            notes = item.get('notes', '')

            for block in _iter_markdown_blocks(content):
                if block[0] == "heading":
                    _, level, title = block
                    doc.add_heading(title, level=min(int(level), 9))
                else:
                    _, para = block
                    doc.add_paragraph(para)
            
            # Handle Notes
            if notes:
                p_note = doc.add_paragraph()
                run = p_note.add_run(f"Note: {notes}")
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100) # Gray
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"[Error] Export to DOCX failed: {e}")
        return False

def export_json_to_pdf(json_path: str, output_path: str) -> bool:
    try:
        data = _load_json_data(json_path)
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph

        fonts = [
            ("SimHei", "C:\\Windows\\Fonts\\simhei.ttf"),
            ("Microsoft YaHei", "C:\\Windows\\Fonts\\msyh.ttf"),
            ("Arial Unicode MS", "C:\\Windows\\Fonts\\ARIALUNI.TTF"),
        ]

        font_name = "Helvetica"
        for _, path in fonts:
            if os.path.exists(path) and path.lower().endswith(".ttf"):
                try:
                    pdfmetrics.registerFont(TTFont("CJKFont", path))
                    font_name = "CJKFont"
                    break
                except Exception as e:
                    print(f"[Warning] Failed to register PDF font: {e}")

        styles = getSampleStyleSheet()
        base_style = ParagraphStyle(
            name="Base",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10.5,
            leading=14,
            spaceAfter=6,
        )
        note_style = ParagraphStyle(
            name="Note",
            parent=base_style,
            fontSize=9,
            leading=12,
            textColor="#666666",
        )
        heading_styles = {}
        for level, size in [(1, 18), (2, 15), (3, 13), (4, 12), (5, 11), (6, 10.5)]:
            heading_styles[level] = ParagraphStyle(
                name=f"H{level}",
                parent=base_style,
                fontSize=size,
                leading=max(int(size * 1.2), 12),
                spaceBefore=10,
                spaceAfter=6,
            )

        doc = SimpleDocTemplate(
            output_path,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=os.path.splitext(os.path.basename(output_path))[0],
        )

        story = []
        for item in data:
            content = item.get("translation", "") or ""
            notes = item.get("notes", "") or ""

            for block in _iter_markdown_blocks(content):
                if block[0] == "heading":
                    _, level, title = block
                    level = min(int(level), 6)
                    story.append(Paragraph(html.escape(title), heading_styles[level]))
                else:
                    _, para = block
                    para_html = html.escape(para).replace("\n", "<br/>")
                    story.append(Paragraph(para_html, base_style))

            if notes:
                story.append(Paragraph(f"Note: {html.escape(notes)}", note_style))

        if not story:
            story.append(Paragraph("(empty)", base_style))

        doc.build(story)
        return True
    except Exception as e:
        print(f"[Error] Export to PDF failed: {e}")
        return False

def export_json_to_epub(json_path: str, output_path: str) -> bool:
    try:
        data = _load_json_data(json_path)
        base_name = os.path.splitext(os.path.basename(json_path))[0].replace('_intermediate', '')
        
        book = epub.EpubBook()
        book.set_identifier(f'id_{base_name}')
        book.set_title(base_name)
        book.set_language('zh')
        book.add_author('Program Translator')
        
        chapters = []
        chapter_count = 1
        current_chapter = epub.EpubHtml(title='Introduction', file_name='intro.xhtml', lang='zh')
        chapter_content = "<h1>Start</h1>"

        for item in data:
            content = item.get('translation', '') or ''
            notes = item.get('notes', '')

            for block in _iter_markdown_blocks(content):
                if block[0] == "heading":
                    _, level, title = block
                    level = int(level)
                    title_html = html.escape(title)
                    if level == 1:
                        if current_chapter:
                            current_chapter.content = chapter_content
                            book.add_item(current_chapter)
                            chapters.append(current_chapter)

                        current_chapter = epub.EpubHtml(
                            title=title,
                            file_name=f'chap_{chapter_count}.xhtml',
                            lang='zh',
                        )
                        chapter_count += 1
                        chapter_content = f"<h1>{title_html}</h1>\n"
                    else:
                        h_level = min(level, 6)
                        chapter_content += f"<h{h_level}>{title_html}</h{h_level}>\n"
                else:
                    _, para = block
                    para_html = html.escape(para).replace("\n", "<br/>")
                    chapter_content += f"<p>{para_html}</p>\n"
            
            if notes:
                chapter_content += f"<p style='color:gray; font-style:italic;'>Note: {html.escape(notes)}</p>\n"
        
        # Save last chapter
        if current_chapter:
            current_chapter.content = chapter_content
            book.add_item(current_chapter)
            chapters.append(current_chapter)
        
        # Define Table of Contents
        book.toc = tuple(chapters)
        
        # Add default NCX and Nav file
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # Define CSS style
        style = 'body { font-family: sans-serif; }'
        nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
        book.add_item(nav_css)
        
        # Basic spine
        book.spine = ['nav'] + chapters
        
        epub.write_epub(output_path, book, {})
        return True
        
    except Exception as e:
        print(f"[Error] Export to EPUB failed: {e}")
        return False
