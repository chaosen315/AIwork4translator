import os
import sys

from docx import Document
from ebooklib import epub
from pypdf import PdfReader


def _print_docx_preview(docx_path: str) -> None:
    print("\n--- DOCX headings (first 30) ---")
    d = Document(docx_path)
    headings = []
    for p in d.paragraphs:
        style_name = p.style.name if p.style else ""
        text = (p.text or "").strip()
        if text and style_name.lower().startswith("heading"):
            headings.append(f"{style_name}: {text}")
        if len(headings) >= 30:
            break
    print("\n".join(headings) if headings else "(no headings found)")

    print("\n--- DOCX sample paragraphs (first 8 non-empty) ---")
    sample = []
    for p in d.paragraphs:
        text = (p.text or "").strip()
        if text:
            sample.append(text)
        if len(sample) >= 8:
            break
    for t in sample:
        print(f"- {t[:200]}")


def _print_epub_preview(epub_path: str) -> None:
    print("\n--- EPUB toc (first 30) ---")
    book = epub.read_epub(epub_path)
    stack = [book.toc]
    titles = []
    while stack:
        node = stack.pop(0)
        if isinstance(node, (list, tuple)):
            stack[0:0] = list(node)
            continue
        title = getattr(node, "title", None)
        if title:
            titles.append(str(title))
    print("\n".join(titles[:30]) if titles else "(toc empty)")


def _print_pdf_preview(pdf_path: str) -> None:
    print("\n--- PDF text extract (page 1, first 800 chars) ---")
    reader = PdfReader(pdf_path)
    if not reader.pages:
        print("(no pages)")
        return
    text = (reader.pages[0].extract_text() or "").replace("\r", "\n")
    print(text[:800])


def main() -> int:
    if len(sys.argv) < 2:
        print("Usage: uv run python scripts/preview_export_outputs.py <base_path_without_ext>")
        return 2

    base = sys.argv[1]
    docx_path = base + ".docx"
    pdf_path = base + ".pdf"
    epub_path = base + ".epub"

    missing = [p for p in [docx_path, pdf_path, epub_path] if not os.path.exists(p)]
    if missing:
        print("Missing files:")
        for p in missing:
            print(f"- {p}")
        return 1

    print(
        "SIZES(bytes):",
        {
            "docx": os.path.getsize(docx_path),
            "pdf": os.path.getsize(pdf_path),
            "epub": os.path.getsize(epub_path),
        },
    )

    _print_docx_preview(docx_path)
    _print_epub_preview(epub_path)
    _print_pdf_preview(pdf_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

